from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
SOURCE_DOCX = Path(r"C:\Users\xsanf\Downloads\Telegram Desktop\HSE_Bachelor_Thesis_LawVoice.docx")
OUTPUT_DOCX = SOURCE_DOCX.with_name("HSE_Bachelor_Thesis_LawVoice_revised_defense.docx")
NOTES_MD = ROOT / "docs" / "defense" / "LawVoice_Thesis_Defense_Fixes.md"
RAG_JSON = ROOT / "docs" / "defense" / "lawvoice_retrieval_benchmark.json"
BACKEND_JSON = ROOT / "docs" / "defense" / "lawvoice_backend_benchmark.json"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    if style:
        new_paragraph.style = style
    return new_paragraph


def find_paragraph(document: Document, predicate) -> Paragraph:
    for paragraph in document.paragraphs:
        if predicate(paragraph.text.strip()):
            return paragraph
    raise ValueError("Paragraph not found")


def paragraph_index(document: Document, target: Paragraph) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph._p is target._p:
            return index
    raise ValueError("Paragraph index not found")


def benchmark_summary_texts() -> tuple[str, str]:
    rag = load_json(RAG_JSON)
    backend = load_json(BACKEND_JSON)

    rag_metrics = rag["metrics"]
    backend_lines = []
    for item in backend["results"]:
        backend_lines.append(
            f"{item['route']}: p50 {item['p50_ms']} ms, p95 {item['p95_ms']} ms, throughput {item['throughput_rps']} requests/s"
        )

    abstract_benchmark = (
        f"In addition, a small offline retrieval benchmark over the shipped internal knowledge corpus achieved "
        f"Hit@1={rag_metrics['hit_at_1']:.2f}, Hit@3={rag_metrics['hit_at_3']:.2f}, "
        f"MRR@3={rag_metrics['mrr_at_3']:.2f} and nDCG@3={rag_metrics['ndcg_at_3']:.3f}, while a synthetic "
        f"in-process backend benchmark reported p95 latency from "
        f"{min(item['p95_ms'] for item in backend['results']):.2f} ms to "
        f"{max(item['p95_ms'] for item in backend['results']):.2f} ms on representative routes with mocked external "
        f"dependencies."
    )
    backend_detail = " ; ".join(backend_lines)
    return abstract_benchmark, backend_detail


def revise_document() -> None:
    document = Document(str(SOURCE_DOCX))
    abstract_benchmark, backend_detail = benchmark_summary_texts()
    rag = load_json(RAG_JSON)
    rag_metrics = rag["metrics"]

    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The proposed solution combines a browser-based WebRTC speech interface")),
        (
            "The proposed solution combines a browser-based WebRTC speech interface, an Express backend, ephemeral "
            "session authentication, a hybrid intent router, a scenario-aware dialog manager, and a pre-generation "
            "retrieval layer that assembles grounded context before the LLM produces an answer. In this architecture, "
            "RAG is not used to justify or rewrite model output after the fact; instead, retrieval narrows the "
            "evidence set that the model may use when generating the response and the subsequent action plan."
        ),
    )
    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("Static analysis of the archive identified 57 relevant source")),
        (
            "The evaluation therefore covers both implementation breadth and operational behavior. Static analysis of "
            "the archive identified 57 relevant source/configuration/documentation files and 11,140 logical lines of "
            "code excluding virtual environments, bundled distribution files and logs. The backend exposes 32 API "
            "routes, and the repository includes 13 automated tests. "
            + abstract_benchmark
            + " These numbers do not claim full end-to-end legal correctness, but they make the prototype evaluation "
              "explicit and measurable."
        ),
    )
    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The aim of this work is to design, implement and evaluate a prototype architecture")),
        (
            "The aim of this work is to design, implement and evaluate a backend-centered prototype of a real-time "
            "retrieval-augmented voice assistant for adolescent legal literacy, with explicit attention to where "
            "retrieval occurs in the pipeline, how safety is enforced, and how the system can be observed and measured."
        ),
    )

    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("Evaluate the prototype through static project analysis")),
        (
            "Evaluate the prototype through static project analysis, a proxy retrieval benchmark, a synthetic backend "
            "benchmark, qualitative scenario analysis and a security/risk matrix."
        ),
    )
    bullet_after = find_paragraph(document, lambda t: t.startswith("Evaluate the prototype through static project analysis"))
    insert_paragraph_after(
        bullet_after,
        (
            "Specify an observability stack based on Langfuse, Prometheus and Grafana, and identify the main "
            "scalability bottlenecks of the current prototype."
        ),
        style="Normal",
    )

    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The personal contribution of this work is the system-level integration")),
        (
            "The personal contribution of this work is the system-level integration and analysis of the LawVoice "
            "prototype: domain framing, requirements, architecture, implementation review, safety design, retrieval "
            "pipeline clarification, evaluation design, observability plan and documentation. The work emphasizes "
            "applied data-science engineering rather than legal doctrinal analysis."
        ),
    )
    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The scope is prototype-level rather than clinical or legal validation")),
        (
            "The scope is a backend-centered prototype evaluation rather than full legal or pedagogical validation. "
            "The work defends the architecture, control flow and measurable prototype behavior, not a claim of legal "
            "authority, institutional confidentiality or production readiness."
        ),
    )

    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The prototype is evaluated against four criteria")),
        (
            "The prototype is evaluated against five criteria: implementation completeness, architectural fit to the "
            "domain, retrieval quality before generation, operational behavior of the backend, and safety/privacy risk "
            "coverage. Retrieval quality is measured with Hit@k, MRR and nDCG on a small benchmark; operational "
            "behavior is measured via route-level latency and throughput in a synthetic benchmark. Because the supplied "
            "archive does not include a production legal corpus or user-study data, these measurements are presented as "
            "prototype-level evidence rather than as final product validation."
        ),
    )

    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The proposed architecture separates the realtime media channel")),
        (
            "The proposed architecture separates the realtime media channel from the domain-control layer and the "
            "knowledge layer. This separation is central: the browser handles audio capture and playback, the backend "
            "controls authentication, safety policy, retrieval and logging, and the model is invoked only after the "
            "relevant context has been retrieved and the prompt has been assembled. The defended sequence is therefore "
            "user query -> intent/risk detection -> retrieval -> grounded prompt assembly -> LLM response generation "
            "-> post-generation validation and delivery."
        ),
    )
    arch_anchor = find_paragraph(document, lambda t: t.startswith("The proposed architecture separates the realtime media channel"))
    insert_paragraph_after(
        arch_anchor,
        (
            "In component terms, the browser communicates with the backend for session creation, CSRF protection and "
            "domain APIs; the backend communicates with PostgreSQL/pgvector for retrieval and with the external model "
            "provider for generation; observability components record traces, metrics and alerts around these "
            "boundaries rather than only around generic application logs."
        ),
        style="Normal",
    )

    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The knowledge module implements document ingestion")),
        (
            "The knowledge module implements document ingestion, text normalization, chunking, token estimation, "
            "embedding creation and storage in PostgreSQL. When pgvector is available, it uses vector similarity with "
            "an HNSW index; otherwise it falls back to Russian full-text search. Retrieval happens before answer "
            "generation: the backend first resolves top-k evidence chunks, then assembles a restricted context and "
            "only then sends the request to the model. This matters methodologically: a RAG system should ground the "
            "model input, not post-process the model output."
        ),
    )
    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The planner receives an objective, context, constraints")),
        (
            "The planner receives an objective, context, constraints, current plan and retrieved knowledge hits. It "
            "asks the model for JSON-only output and then validates the result. Crucially, evidence_ids are filtered "
            "so that the model can cite only chunk IDs that were actually returned by retrieval. Post-generation "
            "validation constrains the response format and references; it is not a substitute for retrieval itself."
        ),
    )
    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("Algorithm 1 - Evidence-bound action planning")),
        (
            "Algorithm 1 - Pre-generation evidence-bound action planning\n"
            "1. Normalize objective, context, constraints and current plan.\n"
            "2. Retrieve top-k knowledge chunks by vector similarity or full-text search.\n"
            "3. Build a restricted prompt from the retrieved chunks only.\n"
            "4. Request a JSON action plan from the model.\n"
            "5. Validate the JSON structure and filter evidence_ids to the retrieved set.\n"
            "6. Return the grounded plan and citations."
        ),
    )

    chapter6_heading = find_paragraph(document, lambda t: t == "6 Evaluation and Results")
    chapter6_index = paragraph_index(document, chapter6_heading)
    cursor = insert_paragraph_after(document.paragraphs[chapter6_index - 1], "5.6 Observability, monitoring and request tracing", style="Heading 2")
    cursor = insert_paragraph_after(
        cursor,
        (
            "If the work is defended as a backend/system prototype, observability must be treated as part of the "
            "architecture rather than optional developer tooling. At the LLM boundary, Langfuse is the most suitable "
            "tracing layer because it records prompts, retrieved context, model outputs, token usage, latency and "
            "failure cases in a form directly relevant to RAG debugging and safety review [10]. In LawVoice, the "
            "natural Langfuse instrumentation points are /api/classify-intent, /api/knowledge/search, /api/action-plan "
            "and realtime session creation."
        ),
        style="Normal",
    )
    cursor = insert_paragraph_after(
        cursor,
        (
            "For infrastructure monitoring, Prometheus and Grafana form the minimal metrics-and-dashboard stack [11,12]. "
            "The backend should export request_count, error_rate, latency histograms, empty_retrieval_rate, "
            "retrieval_latency_ms, action_plan_latency_ms, token_cost_usd, provider_timeout_count and "
            "high_risk_escalation_count. Alerting should trigger on sustained 5xx growth, repeated zero-hit retrieval, "
            "abnormal cost spikes and provider unavailability."
        ),
        style="Normal",
    )
    cursor = insert_paragraph_after(cursor, "5.7 Scalability model and deployment boundaries", style="Heading 2")
    cursor = insert_paragraph_after(
        cursor,
        (
            "The current prototype is structurally scalable at the stateless backend layer but not yet "
            "production-hardened. Session bootstrap, intent classification, retrieval orchestration and action-plan "
            "generation can be replicated behind a reverse proxy because they do not require sticky in-memory state. "
            "The main scaling bottlenecks are PostgreSQL/pgvector query latency, external model latency, and any shared "
            "analytics or cache store."
        ),
        style="Normal",
    )
    insert_paragraph_after(
        cursor,
        (
            "A production decomposition would separate four roles: (1) realtime session gateway, (2) domain "
            "dialog/orchestration service, (3) retrieval service backed by PostgreSQL/pgvector, and (4) "
            "observability/analytics services. This work does not claim that such a deployment has already been "
            "validated at production scale; it identifies the boundaries and failure points that must be addressed next."
        ),
        style="Normal",
    )

    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The evaluation combines artifact inspection and scenario analysis")),
        (
            "The evaluation combines artifact inspection, a proxy retrieval benchmark, a synthetic backend benchmark "
            "and scenario analysis. Static project analysis measures implementation breadth and checks whether the "
            "architecture described in this thesis is reflected in code. The retrieval benchmark measures whether "
            "relevant chunks can be found before generation. The backend benchmark measures route-level latency and "
            "throughput without external-provider noise. Scenario analysis checks whether the assistant has explicit "
            "states and escalation behavior for representative legal-literacy situations."
        ),
    )
    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("This is a pre-defense level evaluation")),
        (
            "This remains a prototype-level evaluation. It does not claim final legal correctness or user "
            "effectiveness. However, unlike a purely descriptive pre-defense draft, the revised evaluation includes "
            "explicit retrieval and backend metrics so that the architectural claims are measurable and falsifiable."
        ),
    )

    eval_anchor = find_paragraph(document, lambda t: t.startswith("This remains a prototype-level evaluation"))
    eval_cursor = insert_paragraph_after(eval_anchor, "6.2 Retrieval benchmark and RAG metrics", style="Heading 2")
    eval_cursor = insert_paragraph_after(
        eval_cursor,
        (
            "RAG must be evaluated before generation, because the function of retrieval is to determine whether the "
            "model receives relevant evidence in the first place. For this reason the key retrieval metrics are Hit@k, "
            "MRR@k and nDCG@k. Hit@k measures whether at least one relevant chunk appears in the first k results. "
            "MRR@k captures how early the first relevant chunk appears. nDCG@k additionally rewards correct ranking "
            "when multiple relevant chunks exist."
        ),
        style="Normal",
    )
    eval_cursor = insert_paragraph_after(
        eval_cursor,
        (
            "The repository snapshot does not include a reproducible PostgreSQL knowledge dump with expert labels, so "
            "a small offline proxy benchmark was built over the shipped internal knowledge corpus using the same "
            "chunking policy as the prototype and a Russian-friendly character n-gram TF-IDF retriever. The corpus "
            f"contained {rag['corpus']['documents']} documents and {rag['corpus']['chunks']} chunks. On "
            f"{rag_metrics['queries']} manually mapped queries, the benchmark produced Hit@1={rag_metrics['hit_at_1']:.2f}, "
            f"Hit@3={rag_metrics['hit_at_3']:.2f}, MRR@3={rag_metrics['mrr_at_3']:.2f} and "
            f"nDCG@3={rag_metrics['ndcg_at_3']:.3f}. These figures should be interpreted as a sanity check for "
            "retrievability of the curated corpus, not as a final legal-quality benchmark."
        ),
        style="Normal",
    )
    eval_cursor = insert_paragraph_after(
        eval_cursor,
        (
            "The most important methodological conclusion is not the absolute value of these proxy metrics but the "
            "correct placement of retrieval in the pipeline. The system first retrieves evidence, then asks the model "
            "to generate a grounded answer or action plan. This corrected description resolves the common mistake of "
            "treating RAG as a post-hoc justification stage."
        ),
        style="Normal",
    )

    eval_cursor = insert_paragraph_after(eval_cursor, "6.3 Synthetic backend performance and operational observations", style="Heading 2")
    eval_cursor = insert_paragraph_after(
        eval_cursor,
        (
            "To assess backend behavior separately from third-party model latency, a synthetic in-process benchmark was "
            "executed with mocked OpenAI/PostgreSQL/Redis dependencies. This benchmark does not represent Internet or "
            "provider delay; it measures the control-plane overhead of the prototype itself."
        ),
        style="Normal",
    )
    eval_cursor = insert_paragraph_after(
        eval_cursor,
        (
            "The results are as follows. "
            + backend_detail
            + " For a prototype, these values indicate that the local backend overhead is small compared with the "
              "expected external LLM latency."
        ),
        style="Normal",
    )
    insert_paragraph_after(
        eval_cursor,
        (
            "A separate observation from the benchmark is that the global rate limiter becomes visible quickly under "
            "repeated protected calls. This is desirable from a protective standpoint, but it also means that a "
            "production version should distinguish anonymous abuse protection from internal service traffic, dashboards "
            "and administrative workflows."
        ),
        style="Normal",
    )

    replace_paragraph_text(find_paragraph(document, lambda t: t == "6.2 Implementation depth results"), "6.4 Implementation depth results")
    replace_paragraph_text(find_paragraph(document, lambda t: t == "6.3 Qualitative scenario evaluation"), "6.5 Qualitative scenario evaluation")
    replace_paragraph_text(find_paragraph(document, lambda t: t == "6.4 Security and risk assessment"), "6.6 Security and privacy assessment")

    security_paragraph = find_paragraph(document, lambda t: t.startswith("Figure 4 maps known LLM application risks"))
    replace_paragraph_text(
        security_paragraph,
        (
            "Figure 4 maps known LLM application risks to controls visible in the project. The strongest controls are "
            "architectural: the browser never receives a standard API key, high-risk routing is implemented in code, "
            "retrieved evidence IDs are constrained, and tool calls are whitelisted. At the same time, the security "
            "claim must be limited. Because the current prototype uses a proprietary external LLM API, user data may "
            "leave the institutional boundary and be processed by provider-managed infrastructure. Therefore the correct "
            "claim is not 'secure by design' in an absolute sense, but 'risk-reduced by architecture and operational controls.'"
        ),
    )
    insert_paragraph_after(
        security_paragraph,
        (
            "From a privacy perspective, the most relevant next step is request-level observability with redaction-aware "
            "tracing rather than generic application logs. Langfuse is a suitable choice because it can record prompts, "
            "retrieved context, outputs and token/cost metadata while supporting selective masking and audit of "
            "high-risk cases."
        ),
        style="Normal",
    )

    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The main finding is that a modern voice assistant")),
        (
            "The main finding is that a modern voice assistant for a high-risk educational domain should be designed as "
            "a backend-controlled socio-technical system, not as a simple wrapper around a realtime model. LawVoice is "
            "therefore best defended as a prototype of a controlled architecture in which retrieval, safety routing, "
            "observability and generation are explicitly separated and measured."
        ),
    )
    replace_paragraph_text(
        find_paragraph(document, lambda t: t.startswith("The prototype shows strong architectural coverage")),
        (
            "The prototype shows broad architectural coverage: realtime WebRTC connection, server-side credentials, "
            "domain prompts, scenario manager, retrieval ingestion and search, evidence-bound planning, analytics and "
            "tests. After the pre-defense corrections, the work also makes explicit what is measured and what is not: "
            "proxy retrieval quality is measured, local backend overhead is measured, whereas full legal correctness, "
            "production privacy guarantees and human educational outcomes remain future work."
        ),
    )

    limitation_bullets = [
        "The evaluation combines static analysis, a proxy retrieval benchmark and a synthetic backend benchmark; no controlled user study was performed in this environment.",
        "The retrieval benchmark uses the shipped internal corpus rather than an external expert-labelled legal benchmark, so it demonstrates retrievability rather than final legal answer quality.",
        "Because the current prototype relies on a proprietary external LLM API, it cannot claim full institutional confidentiality; it can only reduce exposure through architectural controls.",
        "The observability stack is specified at the architectural level, but the archive does not yet ship a production Langfuse/Prometheus/Grafana deployment bundle.",
        "The synthetic backend benchmark isolates control-plane overhead and excludes provider latency, network jitter and real database contention.",
        "The repository contains commerce/product-search functionality inherited from the broader prototype; this should be separated or disabled for a focused LawVoice deployment.",
        "Privacy, data retention and parental/educational consent policies must be defined before real use with minors."
    ]
    limit_anchor = find_paragraph(document, lambda t: t.startswith("The prototype was evaluated primarily through archive inspection"))
    replace_paragraph_text(limit_anchor, limitation_bullets[0])
    following_limitations = [
        find_paragraph(document, lambda t: t.startswith("Legal correctness depends on the quality")),
        find_paragraph(document, lambda t: t.startswith("Prompt injection cannot be fully prevented")),
        find_paragraph(document, lambda t: t.startswith("The current tests cover important logic")),
        find_paragraph(document, lambda t: t.startswith("The repository contains commerce/product-search functionality")),
        find_paragraph(document, lambda t: t.startswith("Privacy, data retention and parental/educational consent policies")),
    ]
    replacement_texts = limitation_bullets[1:6]
    for paragraph, text in zip(following_limitations, replacement_texts, strict=True):
        replace_paragraph_text(paragraph, text)
    insert_paragraph_after(following_limitations[-1], limitation_bullets[6], style="Normal")

    future_bullets = [
        "Create a legally reviewed knowledge base with source metadata, jurisdiction, validity dates and age-appropriate explanations.",
        "Integrate Langfuse tracing with prompt redaction, retrieved-context logging and review workflows for high-risk cases.",
        "Export Prometheus metrics, create Grafana dashboards and define alert rules for provider timeouts, repeated zero-hit retrieval and abnormal token-cost spikes.",
        "Add adversarial evaluation: prompt injection, indirect injection in documents, jailbreaks, system-prompt leakage and RAG poisoning tests.",
        "Run real end-to-end load tests with database-backed deployment, external model latency and concurrent session traffic.",
        "Run expert evaluation with lawyers, teachers and child-safety specialists using a rubric for correctness, safety, clarity and empathy.",
        "Introduce human-in-the-loop escalation for high-risk cases and a clear disclaimer flow before advice-like content."
    ]
    future_anchor = find_paragraph(document, lambda t: t.startswith("Create a legally reviewed knowledge base"))
    replace_paragraph_text(future_anchor, future_bullets[0])
    future_others = [
        find_paragraph(document, lambda t: t.startswith("Add adversarial evaluation")),
        find_paragraph(document, lambda t: t.startswith("Measure end-to-end latency")),
        find_paragraph(document, lambda t: t.startswith("Run expert evaluation")),
        find_paragraph(document, lambda t: t.startswith("Introduce human-in-the-loop escalation")),
        find_paragraph(document, lambda t: t.startswith("Refactor the prototype into separate services")),
    ]
    future_replacements = future_bullets[1:6]
    for paragraph, text in zip(future_others, future_replacements, strict=True):
        replace_paragraph_text(paragraph, text)
    insert_paragraph_after(future_others[-1], future_bullets[6], style="Normal")

    conclusion_bullets = [
        "We formulated the problem, aim and requirements for a voice assistant that must combine natural interaction with factual grounding, safety, observability and privacy-aware backend control.",
        "We proposed a modular architecture in which retrieval is placed before generation, while post-generation validation constrains format and citations rather than replacing RAG.",
        "We analysed the supplied project archive and found 57 relevant files, 11,140 logical LOC, 32 backend routes and 13 automated tests.",
        f"We reported a proxy retrieval benchmark over the shipped internal corpus with Hit@1={rag_metrics['hit_at_1']:.2f}, Hit@3={rag_metrics['hit_at_3']:.2f}, MRR@3={rag_metrics['mrr_at_3']:.2f} and nDCG@3={rag_metrics['ndcg_at_3']:.3f}.",
        "We reported a synthetic backend benchmark showing low local control-plane latency on representative routes and documented the implications of the global rate limiter.",
        "We produced a risk matrix and an observability/scalability plan that limits security claims to risk reduction rather than absolute confidentiality."
    ]
    concl_anchor = find_paragraph(document, lambda t: t.startswith("We formulated the problem, aim and requirements"))
    replace_paragraph_text(concl_anchor, conclusion_bullets[0])
    concl_others = [
        find_paragraph(document, lambda t: t.startswith("We proposed a modular architecture based on WebRTC")),
        find_paragraph(document, lambda t: t.startswith("We analysed the supplied project archive")),
        find_paragraph(document, lambda t: t.startswith("We showed that high-risk legal-literacy scenarios")),
        find_paragraph(document, lambda t: t.startswith("We produced a risk matrix mapping prompt injection")),
        find_paragraph(document, lambda t: t.startswith("We identified limitations and a roadmap")),
    ]
    for paragraph, text in zip(concl_others, conclusion_bullets[1:], strict=True):
        replace_paragraph_text(paragraph, text)

    bibliography_anchor = find_paragraph(document, lambda t: t.startswith("[9] Project archive"))
    biblio_cursor = insert_paragraph_after(
        bibliography_anchor,
        "[10] Langfuse Documentation. Observability and tracing for LLM applications. URL: https://langfuse.com/docs (accessed 07.05.2026).",
        style="Normal",
    )
    biblio_cursor = insert_paragraph_after(
        biblio_cursor,
        "[11] Prometheus Documentation. Overview and instrumentation concepts. URL: https://prometheus.io/docs/ (accessed 07.05.2026).",
        style="Normal",
    )
    insert_paragraph_after(
        biblio_cursor,
        "[12] Grafana Documentation. Dashboards and observability stack overview. URL: https://grafana.com/docs/ (accessed 07.05.2026).",
        style="Normal",
    )

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_DOCX))

    notes = f"""# LawVoice Thesis Defense Fixes

## Output

- Revised thesis DOCX: `{OUTPUT_DOCX}`
- Retrieval benchmark JSON: `{RAG_JSON}`
- Backend benchmark JSON: `{BACKEND_JSON}`

## Main thesis corrections

1. The goal was narrowed to a backend-centered prototype of a real-time RAG voice assistant, with explicit emphasis on retrieval placement, safety enforcement and measurability.
2. The RAG description was corrected so that retrieval clearly happens before generation: user query -> intent/risk detection -> retrieval -> grounded prompt assembly -> LLM response -> post-generation validation.
3. The work now includes explicit retrieval metrics: Hit@1={rag_metrics['hit_at_1']:.2f}, Hit@3={rag_metrics['hit_at_3']:.2f}, MRR@3={rag_metrics['mrr_at_3']:.2f}, nDCG@3={rag_metrics['ndcg_at_3']:.3f}.
4. The evaluation chapter now includes a synthetic backend benchmark with route-level latency and throughput numbers.
5. Security claims were narrowed: the revised text states clearly that a proprietary external LLM prevents any claim of full confidentiality, so the correct claim is risk reduction rather than absolute security.
6. Observability was promoted to an architectural requirement. The revised thesis explicitly recommends Langfuse for LLM request tracing and Prometheus/Grafana for metrics, dashboards and alerts.
7. The scalability section now identifies the stateless backend layer, the PostgreSQL/pgvector retrieval layer and provider latency as the main performance boundaries.

## Backend benchmark snapshot

- {backend_detail}

## Defense positioning

- Defend the work as a backend/system prototype, not as a fully validated legal product.
- When asked about RAG, say that retrieval is evaluated before generation and that post-generation validation only constrains format and citations.
- When asked about security, say that the prototype reduces risk but does not provide absolute confidentiality because the provider-side LLM remains external.
- When asked about scaling, say that the current artifact is single-node and prototype-level, but the thesis now identifies the concrete bottlenecks and the target service decomposition.
"""
    NOTES_MD.write_text(notes, encoding="utf-8")


if __name__ == "__main__":
    revise_document()
    print(OUTPUT_DOCX)
    print(NOTES_MD)
